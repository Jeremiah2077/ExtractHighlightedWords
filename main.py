import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import unicodedata

def get_highlight_color(file_path, word):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.lower().find(word.lower()) != -1 and run.font.highlight_color is not None:
                return run.font.highlight_color
    return None

def is_chinese(char):
    return 'CJK UNIFIED' in unicodedata.name(char)

def find_highlighted_words_with_color(file_path, color):
    highlighted_words = []
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == color:
                words = run.text.split()
                words = [word for word in words if not any(is_chinese(c) for c in word)]
                highlighted_words.extend(words)
    return highlighted_words

def main(file_path, demo_word):
    color = get_highlight_color(file_path, demo_word)
    if color is not None:
        highlighted_words = find_highlighted_words_with_color(file_path, color)
        return highlighted_words
    else:
        return []

def select_file():
    file_path = filedialog.askopenfilename()
    file_path_entry.delete(0, tk.END)
    file_path_entry.insert(0, file_path)

def extract_words():
    file_path = file_path_entry.get()
    demo_word = demo_word_entry.get()
    highlighted_words = main(file_path, demo_word)
    result_text.delete(1.0, tk.END)
    result_text.insert(1.0, ', '.join(highlighted_words))
    count_label.config(text="Count: " + str(len(highlighted_words)))

def copy_result():
    root.clipboard_clear()
    root.clipboard_append(result_text.get(1.0, tk.END))

root = tk.Tk()
root.title("Extract Highlighted Words")

tk.Label(root, text="File Path:").grid(row=0, column=0, sticky='w')
file_path_entry = tk.Entry(root, width=40)
file_path_entry.grid(row=0, column=1, sticky='w')
tk.Button(root, text="Select File", command=select_file).grid(row=0, column=2, sticky='w')

tk.Label(root, text="Demo Word:").grid(row=1, column=0, sticky='w')
demo_word_entry = tk.Entry(root, width=40)
demo_word_entry.grid(row=1, column=1, sticky='w')

tk.Button(root, text="Extract Words", command=extract_words).grid(row=1, column=2, sticky='w')

result_text = tk.Text(root, width=50, height=10)
result_text.grid(row=2, column=0, columnspan=3, pady=10)

tk.Button(root, text="Copy Result", command=copy_result).grid(row=3, column=1)

count_label = tk.Label(root, text="Count: 0", fg="green")
count_label.grid(row=3, column=0)

root.mainloop()
